import fitz  # PyMuPDF
import docx
import tempfile
import os
from uuid import uuid4
import logging

logger = logging.getLogger(__name__)


async def extract_text_from_pdf(file_bytes: bytes) -> dict:
    """
    Extracts text from a PDF file.
    
    Parameters:
        file_bytes (bytes): The PDF file contents as bytes
        
    Returns:
        dict: Contains 'text', 'page_count', 'char_count'
        
    Raises:
        ValueError: If PDF has no extractable text or other extraction error
    """
    logger.debug("Starting PDF text extraction")
    
    tmp_path = None
    try:
        # Write to temp file
        tmp_path = os.path.join(tempfile.gettempdir(), f"{uuid4()}.pdf")
        with open(tmp_path, "wb") as f:
            f.write(file_bytes)
        
        # Open and extract
        doc = fitz.open(tmp_path)
        text = ""
        
        # Get page count before closing
        page_count = doc.page_count
        
        for page in doc:
            text += page.get_text("text") + "\n\n"
        
        doc.close()
        
        text = text.strip()
        
        if not text:
            raise ValueError("PDF has no extractable text. Please use /extract/docx-to-text for scanned PDFs.")
        
        result = {
            "text": text,
            "page_count": page_count,
            "char_count": len(text)
        }
        
        logger.debug(f"PDF extraction complete: {result['page_count']} pages, {result['char_count']} chars")
        return result
        
    except Exception as e:
        logger.error(f"PDF extraction error: {str(e)}")
        raise ValueError(f"Error processing PDF file: {str(e)}")
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.remove(tmp_path)


async def extract_text_from_docx(file_bytes: bytes) -> dict:
    """
    Extracts raw text from a .docx Word document file.
    
    Uses python-docx to read all paragraphs from the document.
    Also reads text from tables inside the document.
    
    Parameters:
        file_bytes: Raw bytes of the .docx file
        
    Returns:
        dict with keys:
            - text: str — full extracted text
            - paragraph_count: int — number of paragraphs found
            - char_count: int — total character count
            
    Raises:
        ValueError: if the file is not a valid .docx file
        ValueError: if the document has no extractable text
    """
    logger.debug("Starting DOCX text extraction")
    
    tmp_path = None
    try:
        # Write to temp file
        tmp_path = os.path.join(tempfile.gettempdir(), f"{uuid4()}.docx")
        with open(tmp_path, "wb") as f:
            f.write(file_bytes)
        
        # Open with python-docx
        doc = docx.Document(tmp_path)
        
        # Extract text from paragraphs
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text.strip())
        
        # Join with newlines
        text = "\n".join(full_text).strip()
        
        if not text:
            raise ValueError("Document has no extractable text.")
        
        result = {
            "text": text,
            "paragraph_count": len(doc.paragraphs),
            "char_count": len(text)
        }
        
        logger.debug(f"DOCX extraction complete: {result['paragraph_count']} paragraphs, {result['char_count']} chars")
        return result
        
    except docx.oxml.parse.OxmlElement as e:
        raise ValueError(f"Invalid DOCX file: {str(e)}")
    except Exception as e:
        if "no extractable text" in str(e).lower():
            raise ValueError(str(e))
        raise ValueError(f"DOCX processing error: {str(e)}")
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.remove(tmp_path)

